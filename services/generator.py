from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.enums import TA_LEFT
from docx import Document
from docx.shared import Inches
import io
import os
import re
from PIL import Image as PILImage
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from PIL import ImageChops

def _get_cropped_image_path(original_path: str):
    try:
        img = PILImage.open(original_path).convert("L")  # grayscale
        bg = PILImage.new("L", img.size, 255)            # white background
        diff = ImageChops.difference(img, bg)
        bbox = diff.getbbox()

        if bbox:
            img = img.crop(bbox)

        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format="PNG")
        img_byte_arr.seek(0)
        return img_byte_arr
    except Exception as e:
        print(f"Error cropping image: {e}")
        return None

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SIGNATURE_PATH = os.path.join(BASE_DIR, "static", "signature.png")

print("Using signature:", SIGNATURE_PATH)
print("Signature exists:", os.path.exists(SIGNATURE_PATH))
def _safe_paragraph(text, style):
    from reportlab.platypus import Paragraph
    try:
        return Paragraph(text, style)
    except Exception as e:
        # If parsing fails (e.g. unbalanced tags), try stripping tags
        fallback = text.replace('<b>', '').replace('</b>', '').replace('<u>', '').replace('</u>', '')
        try:
            return Paragraph(fallback, style)
        except Exception:
            # Last resort: escape everything
            fallback = fallback.replace('<', '&lt;').replace('>', '&gt;')
            return Paragraph(fallback, style)

def create_filled_pdf(content_text: str) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    # Create a custom style for form content (Monospaced for alignment)
    form_style = ParagraphStyle(
        'FormStyle',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=9,
        leading=12,  # Line spacing
        spaceAfter=6, # Paragraph spacing
        alignment=TA_LEFT
    )
    
    story = []
    
    # Process text to create paragraphs
    # We treat double newlines as new paragraphs
    # We treat single newlines as line breaks
    
    paragraphs = content_text.split('\n')
    
    for para in paragraphs:
        if not para.strip():
            story.append(Spacer(1, 12)) # Add space for empty lines
        else:
            # Protect valid tags using placeholders (Handle spacing and case)
            clean_text = re.sub(r'<\s*u\s*>', '[[[U_TAG]]]', para, flags=re.IGNORECASE)
            clean_text = re.sub(r'<\s*/\s*u\s*>', '[[[END_U_TAG]]]', clean_text, flags=re.IGNORECASE)
            clean_text = re.sub(r'<\s*b\s*>', '[[[B_TAG]]]', clean_text, flags=re.IGNORECASE)
            clean_text = re.sub(r'<\s*/\s*b\s*>', '[[[END_B_TAG]]]', clean_text, flags=re.IGNORECASE)
            
            # Now safely escape all remaining HTML characters
            clean_text = clean_text.replace('&', '&amp;')
            clean_text = clean_text.replace('<', '&lt;').replace('>', '&gt;')
            
            # Restore tags from placeholders
            clean_text = clean_text.replace('[[[U_TAG]]]', '<u>')
            clean_text = clean_text.replace('[[[END_U_TAG]]]', '</u>')
            clean_text = clean_text.replace('[[[B_TAG]]]', '<b>')
            clean_text = clean_text.replace('[[[END_B_TAG]]]', '</b>')

            # Preserve multiple spaces for alignment
            # We replace 2 spaces with 1 nbsp + 1 space to allow some breaking, or just all nbsp
            clean_text = re.sub(r' {2,}', lambda m: '&nbsp;' * len(m.group(0)), clean_text)
            
            # Check for signature placeholder
            # Use regex to be robust against "Signature:" vs "Signature :" vs "Signature   :"
            sig_match = re.search(r"Signature\s*:", para, re.IGNORECASE)
            
            if sig_match and os.path.exists(SIGNATURE_PATH):
                try:
                    # Add "Signature:"
                    story.append(Paragraph(sig_match.group(0), form_style))

                    # Add signature image
                    cropped_img_data = _get_cropped_image_path(SIGNATURE_PATH)
                    if cropped_img_data:
                        im = Image(cropped_img_data)
                    else:
                        im = Image(SIGNATURE_PATH)
                    
                    # Keep aspect ratio
                    im.drawHeight = 40
                    im.drawWidth = im.imageWidth * im.drawHeight / im.imageHeight
                    im.hAlign = 'LEFT'
                    story.append(im)

                except Exception as e:
                    print(f"Error loading signature image: {e}")
                    # Fallback
                    p = _safe_paragraph(clean_text, form_style)
                    story.append(p)
            else:
                p = _safe_paragraph(clean_text, form_style) 
                story.append(p)

    doc.build(story)
    buffer.seek(0)
    return buffer

def create_filled_docx(content_text: str) -> io.BytesIO:
    doc = Document()
    
    # We need to process the text line by line
    for line in content_text.split('\n'):
        p = doc.add_paragraph()
        
        # Check for signature placeholder
        sig_match = re.search(r"Signature\s*:", line, re.IGNORECASE)
        if sig_match and os.path.exists(SIGNATURE_PATH):
             try:
                 # Add "Signature:" text
                 # Preserve the spacing found in regex match
                 run = p.add_run(sig_match.group(0))
                 run.font.name = 'Courier New'
                 from docx.shared import Pt
                 run.font.size = Pt(10)
                 
                 # Add signal to break line? No, just add picture
                 run = p.add_run()
                 run.add_break()
                 
                 cropped_img_data = _get_cropped_image_path(SIGNATURE_PATH)
                 if cropped_img_data:
                     run.add_picture(cropped_img_data, width=Inches(2.0))
                 else:
                     run.add_picture(SIGNATURE_PATH, width=Inches(2.0))
                     
                 continue # Skip processing the rest of this line (which has "NOT PROVIDED")
             except Exception as e:
                 print(f"Error loading signature image: {e}")
                 # Fallback to normal processing
        
        # Split by tags using regex, keeping the delimiters
        # This will separate text into: ["Start ", "<u>", "Underlined", "</u>", " End"]
        parts = re.split(r'(</?[ubUB]>)', line)
        
        # State tracking
        bold = False
        underline = False
        
        for part in parts:
            if not part:
                continue
                
            lower_part = part.lower()
            
            if lower_part == '<b>':
                bold = True
                continue
            elif lower_part == '</b>':
                bold = False
                continue
            elif lower_part == '<u>':
                underline = True
                continue
            elif lower_part == '</u>':
                underline = False
                continue
                
            # It's text content
            run = p.add_run(part)
            run.bold = bold
            run.underline = underline
            
            # Apply Monospace Font for alignment
            run.font.name = 'Courier New'
            # Also set size if needed, default is usually 11, let's keep it standard or 10
            from docx.shared import Pt
            run.font.size = Pt(10)
            
            # Handle multiple spaces for alignment if needed
            # python-docx handles standard spaces well, but for extra width we might need something
            # sticking to standard text behavior for now as DOCX is reflowable
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
